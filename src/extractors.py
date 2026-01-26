#!/usr/bin/env python3
"""
CONTENT EXTRACTORS
==================
Better extraction for web pages, PDFs, and documents.
"""

import os
import re
import logging
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass
class ExtractedContent:
    """Extracted content from a source."""
    text: str
    title: Optional[str] = None
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


# =============================================================================
# WEB EXTRACTION (trafilatura)
# =============================================================================

def extract_web_content(html: str, url: Optional[str] = None) -> ExtractedContent:
    """
    Extract readable content from HTML using trafilatura.
    Falls back to basic extraction if trafilatura unavailable.
    
    Args:
        html: Raw HTML content
        url: Optional URL for metadata
        
    Returns:
        ExtractedContent with clean text
    """
    try:
        import trafilatura
        
        # Extract with trafilatura (best quality)
        text = trafilatura.extract(
            html,
            include_comments=False,
            include_tables=True,
            include_links=False,
            include_images=False,
            favor_precision=True,
            deduplicate=True
        )
        
        if text:
            # Try to get title
            title = None
            metadata = trafilatura.extract_metadata(html)
            if metadata:
                title = metadata.title
            
            return ExtractedContent(
                text=text,
                title=title,
                metadata={"extractor": "trafilatura", "url": url}
            )
    except ImportError:
        logger.warning("trafilatura not installed, using fallback extraction")
    except Exception as e:
        logger.warning(f"trafilatura extraction failed: {e}")
    
    # Fallback: basic regex extraction
    return _basic_html_extraction(html, url)


def _basic_html_extraction(html: str, url: Optional[str] = None) -> ExtractedContent:
    """Basic HTML text extraction using regex."""
    # Remove script and style
    text = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<head[^>]*>.*?</head>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<nav[^>]*>.*?</nav>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<footer[^>]*>.*?</footer>', '', text, flags=re.DOTALL | re.IGNORECASE)
    
    # Extract title
    title = None
    title_match = re.search(r'<title[^>]*>(.*?)</title>', html, re.DOTALL | re.IGNORECASE)
    if title_match:
        import html as html_module
        title = html_module.unescape(title_match.group(1).strip())
    
    # Replace block elements with newlines
    text = re.sub(r'<(p|div|h[1-6]|li|tr|br)[^>]*>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</(p|div|h[1-6]|li|tr)>', '\n', text, flags=re.IGNORECASE)
    
    # Remove remaining tags
    text = re.sub(r'<[^>]+>', '', text)
    
    # Decode entities
    import html as html_module
    text = html_module.unescape(text)
    
    # Clean whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = re.sub(r' +', ' ', text)
    text = '\n'.join(line.strip() for line in text.split('\n'))
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return ExtractedContent(
        text=text.strip(),
        title=title,
        metadata={"extractor": "basic", "url": url}
    )


# =============================================================================
# PDF EXTRACTION (PyMuPDF/fitz)
# =============================================================================

def extract_pdf_content(pdf_path: str) -> ExtractedContent:
    """
    Extract text from PDF using PyMuPDF.
    
    Args:
        pdf_path: Path to PDF file
        
    Returns:
        ExtractedContent with extracted text
    """
    path = Path(pdf_path)
    if not path.exists():
        raise FileNotFoundError(f"PDF not found: {pdf_path}")
    
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(str(path))
        
        # Extract metadata
        metadata = {
            "extractor": "pymupdf",
            "page_count": len(doc),
            "title": doc.metadata.get("title"),
            "author": doc.metadata.get("author"),
            "subject": doc.metadata.get("subject"),
        }
        
        # Extract text from all pages
        texts = []
        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                texts.append(f"--- Page {page_num + 1} ---\n{text}")
        
        doc.close()
        
        full_text = "\n\n".join(texts)
        
        return ExtractedContent(
            text=full_text,
            title=metadata.get("title") or path.stem,
            metadata=metadata
        )
        
    except ImportError:
        raise ImportError("PyMuPDF required: pip install PyMuPDF")


def extract_pdf_with_structure(pdf_path: str) -> Tuple[str, List[Dict]]:
    """
    Extract PDF with structural information (headings, paragraphs).
    
    Returns:
        Tuple of (full_text, list of section dicts)
    """
    try:
        import fitz
        
        doc = fitz.open(str(pdf_path))
        sections = []
        current_section = {"title": "Introduction", "content": []}
        
        for page_num, page in enumerate(doc):
            # Get text blocks with formatting
            blocks = page.get_text("dict")["blocks"]
            
            for block in blocks:
                if "lines" not in block:
                    continue
                
                for line in block["lines"]:
                    text = "".join(span["text"] for span in line["spans"]).strip()
                    if not text:
                        continue
                    
                    # Check if this looks like a heading (larger font, short)
                    avg_size = sum(span["size"] for span in line["spans"]) / len(line["spans"])
                    is_heading = avg_size > 12 and len(text) < 100 and not text.endswith('.')
                    
                    if is_heading:
                        # Save current section and start new one
                        if current_section["content"]:
                            sections.append(current_section)
                        current_section = {"title": text, "content": [], "page": page_num + 1}
                    else:
                        current_section["content"].append(text)
        
        # Save last section
        if current_section["content"]:
            sections.append(current_section)
        
        doc.close()
        
        # Build full text
        full_text = "\n\n".join(
            f"## {s['title']}\n" + "\n".join(s["content"])
            for s in sections
        )
        
        return full_text, sections
        
    except ImportError:
        raise ImportError("PyMuPDF required: pip install PyMuPDF")


# =============================================================================
# DOCX EXTRACTION (python-docx)
# =============================================================================

def extract_docx_content(docx_path: str) -> ExtractedContent:
    """
    Extract text from Word document.
    
    Args:
        docx_path: Path to .docx file
        
    Returns:
        ExtractedContent with extracted text
    """
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {docx_path}")
    
    try:
        from docx import Document
        
        doc = Document(str(path))
        
        # Extract core properties
        metadata = {
            "extractor": "python-docx",
            "paragraph_count": len(doc.paragraphs),
        }
        
        if doc.core_properties:
            metadata["title"] = doc.core_properties.title
            metadata["author"] = doc.core_properties.author
            metadata["subject"] = doc.core_properties.subject
        
        # Extract text with structure
        texts = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check for heading styles
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading", "").strip() or "1"
                prefix = "#" * int(level) if level.isdigit() else "##"
                texts.append(f"\n{prefix} {text}\n")
            else:
                texts.append(text)
        
        # Also extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                table_text.append(row_text)
            if table_text:
                texts.append("\n" + "\n".join(table_text) + "\n")
        
        full_text = "\n\n".join(texts)
        
        return ExtractedContent(
            text=full_text,
            title=metadata.get("title") or path.stem,
            metadata=metadata
        )
        
    except ImportError:
        raise ImportError("python-docx required: pip install python-docx")


# =============================================================================
# SYNTAX-AWARE CODE CHUNKING
# =============================================================================

def chunk_python_code(code: str, max_chunk_size: int = 1500) -> List[Dict[str, Any]]:
    """
    Chunk Python code using AST for syntax-aware splitting.
    
    Args:
        code: Python source code
        max_chunk_size: Maximum characters per chunk
        
    Returns:
        List of chunks with metadata
    """
    import ast
    
    chunks = []
    
    try:
        tree = ast.parse(code)
        
        # Extract top-level definitions
        for node in ast.iter_child_nodes(tree):
            if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
                # Function definition
                start = node.lineno - 1
                end = node.end_lineno
                lines = code.split('\n')[start:end]
                chunk_text = '\n'.join(lines)
                
                chunks.append({
                    "type": "function",
                    "name": node.name,
                    "text": chunk_text[:max_chunk_size],
                    "lineno": node.lineno,
                    "docstring": ast.get_docstring(node),
                })
                
            elif isinstance(node, ast.ClassDef):
                # Class definition
                start = node.lineno - 1
                end = node.end_lineno
                lines = code.split('\n')[start:end]
                chunk_text = '\n'.join(lines)
                
                # Get method names
                methods = [n.name for n in ast.iter_child_nodes(node) 
                          if isinstance(n, (ast.FunctionDef, ast.AsyncFunctionDef))]
                
                chunks.append({
                    "type": "class",
                    "name": node.name,
                    "text": chunk_text[:max_chunk_size],
                    "lineno": node.lineno,
                    "docstring": ast.get_docstring(node),
                    "methods": methods,
                })
                
            elif isinstance(node, ast.Import):
                # Import statements
                names = [alias.name for alias in node.names]
                chunks.append({
                    "type": "import",
                    "name": ", ".join(names),
                    "text": f"import {', '.join(names)}",
                    "lineno": node.lineno,
                })
                
            elif isinstance(node, ast.ImportFrom):
                # From imports
                names = [alias.name for alias in node.names]
                chunks.append({
                    "type": "import_from",
                    "name": f"{node.module}.{names[0]}" if node.module else names[0],
                    "text": f"from {node.module} import {', '.join(names)}",
                    "lineno": node.lineno,
                })
        
        # If no chunks extracted, fall back to line-based chunking
        if not chunks:
            return _fallback_code_chunking(code, max_chunk_size)
        
        return chunks
        
    except SyntaxError as e:
        logger.warning(f"Python syntax error, falling back to line chunking: {e}")
        return _fallback_code_chunking(code, max_chunk_size)


def chunk_javascript_code(code: str, max_chunk_size: int = 1500) -> List[Dict[str, Any]]:
    """
    Chunk JavaScript/TypeScript code using regex patterns.
    
    Args:
        code: JS/TS source code
        max_chunk_size: Maximum characters per chunk
        
    Returns:
        List of chunks with metadata
    """
    chunks = []
    
    # Patterns for JS/TS constructs
    patterns = [
        # Function declarations
        (r'(?:export\s+)?(?:async\s+)?function\s+(\w+)\s*\([^)]*\)\s*\{', 'function'),
        # Arrow functions assigned to const/let/var
        (r'(?:export\s+)?(?:const|let|var)\s+(\w+)\s*=\s*(?:async\s+)?\([^)]*\)\s*=>', 'arrow_function'),
        # Class declarations
        (r'(?:export\s+)?class\s+(\w+)(?:\s+extends\s+\w+)?\s*\{', 'class'),
        # Interface declarations (TypeScript)
        (r'(?:export\s+)?interface\s+(\w+)(?:\s+extends\s+[\w,\s]+)?\s*\{', 'interface'),
        # Type declarations (TypeScript)
        (r'(?:export\s+)?type\s+(\w+)\s*=', 'type'),
    ]
    
    for pattern, chunk_type in patterns:
        for match in re.finditer(pattern, code):
            name = match.group(1)
            start = match.start()
            
            # Find matching closing brace
            brace_count = 0
            end = start
            in_string = False
            string_char = None
            
            for i, char in enumerate(code[start:], start):
                if char in '"\'`' and not in_string:
                    in_string = True
                    string_char = char
                elif char == string_char and in_string:
                    in_string = False
                elif not in_string:
                    if char == '{':
                        brace_count += 1
                    elif char == '}':
                        brace_count -= 1
                        if brace_count == 0:
                            end = i + 1
                            break
            
            chunk_text = code[start:end]
            if len(chunk_text) > max_chunk_size:
                chunk_text = chunk_text[:max_chunk_size] + "\n// ... truncated"
            
            chunks.append({
                "type": chunk_type,
                "name": name,
                "text": chunk_text,
                "start": start,
            })
    
    if not chunks:
        return _fallback_code_chunking(code, max_chunk_size)
    
    return chunks


def _fallback_code_chunking(code: str, max_chunk_size: int = 1500) -> List[Dict[str, Any]]:
    """Fallback line-based code chunking."""
    lines = code.split('\n')
    chunks = []
    current_chunk = []
    current_size = 0
    
    for i, line in enumerate(lines):
        line_size = len(line) + 1
        
        if current_size + line_size > max_chunk_size and current_chunk:
            chunks.append({
                "type": "block",
                "name": f"lines_{i - len(current_chunk) + 1}-{i}",
                "text": '\n'.join(current_chunk),
                "lineno": i - len(current_chunk) + 1,
            })
            current_chunk = []
            current_size = 0
        
        current_chunk.append(line)
        current_size += line_size
    
    if current_chunk:
        chunks.append({
            "type": "block",
            "name": f"lines_{len(lines) - len(current_chunk) + 1}-{len(lines)}",
            "text": '\n'.join(current_chunk),
            "lineno": len(lines) - len(current_chunk) + 1,
        })
    
    return chunks


def chunk_code_file(filepath: str, max_chunk_size: int = 1500) -> List[Dict[str, Any]]:
    """
    Chunk a code file based on its extension.
    
    Args:
        filepath: Path to code file
        max_chunk_size: Maximum characters per chunk
        
    Returns:
        List of chunks with metadata
    """
    path = Path(filepath)
    ext = path.suffix.lower()
    
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        code = f.read()
    
    if ext == '.py':
        chunks = chunk_python_code(code, max_chunk_size)
    elif ext in {'.js', '.jsx', '.ts', '.tsx', '.mjs'}:
        chunks = chunk_javascript_code(code, max_chunk_size)
    else:
        chunks = _fallback_code_chunking(code, max_chunk_size)
    
    # Add file info to chunks
    for chunk in chunks:
        chunk["file"] = str(path.name)
    
    return chunks


# =============================================================================
# IMAGE DESCRIPTION (Vision API)
# =============================================================================

def describe_image(image_path: str, api_key: Optional[str] = None) -> str:
    """
    Get description of an image using vision model.
    
    Args:
        image_path: Path to image file
        api_key: OpenRouter/OpenAI API key
        
    Returns:
        Text description of the image
    """
    import base64
    
    path = Path(image_path)
    if not path.exists():
        raise FileNotFoundError(f"Image not found: {image_path}")
    
    # Read and encode image
    with open(path, 'rb') as f:
        image_data = base64.b64encode(f.read()).decode('utf-8')
    
    # Determine media type
    ext = path.suffix.lower()
    media_types = {
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png',
        '.gif': 'image/gif',
        '.webp': 'image/webp',
    }
    media_type = media_types.get(ext, 'image/jpeg')
    
    # Get API key
    if api_key is None:
        api_key = os.getenv('OPENROUTER_API_KEY') or os.getenv('OPENAI_API_KEY')
    
    if not api_key:
        raise ValueError("No API key provided for vision model")
    
    try:
        from openai import OpenAI
        
        client = OpenAI(
            api_key=api_key,
            base_url="https://openrouter.ai/api/v1"
        )
        
        response = client.chat.completions.create(
            model="openai/gpt-4o-mini",  # Vision-capable model
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Describe this image in detail. Include: main subjects, actions, colors, text visible, overall mood/context."
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{media_type};base64,{image_data}"
                            }
                        }
                    ]
                }
            ],
            max_tokens=500
        )
        
        return response.choices[0].message.content
        
    except ImportError:
        raise ImportError("openai package required: pip install openai")


def describe_image_url(image_url: str, api_key: Optional[str] = None) -> str:
    """
    Get description of an image from URL.
    
    Args:
        image_url: URL of the image
        api_key: OpenRouter/OpenAI API key
        
    Returns:
        Text description of the image
    """
    if api_key is None:
        api_key = os.getenv('OPENROUTER_API_KEY') or os.getenv('OPENAI_API_KEY')
    
    if not api_key:
        raise ValueError("No API key provided for vision model")
    
    try:
        from openai import OpenAI
        
        client = OpenAI(
            api_key=api_key,
            base_url="https://openrouter.ai/api/v1"
        )
        
        response = client.chat.completions.create(
            model="openai/gpt-4o-mini",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Describe this image in detail. Include: main subjects, actions, colors, text visible, overall mood/context."
                        },
                        {
                            "type": "image_url",
                            "image_url": {"url": image_url}
                        }
                    ]
                }
            ],
            max_tokens=500
        )
        
        return response.choices[0].message.content
        
    except ImportError:
        raise ImportError("openai package required: pip install openai")


if __name__ == "__main__":
    print("Content Extractors Demo")
    print("=" * 50)
    
    # Test HTML extraction
    html = """
    <html>
    <head><title>Test Page</title></head>
    <body>
    <nav>Navigation here</nav>
    <h1>Main Title</h1>
    <p>This is the main content of the page.</p>
    <p>It has multiple paragraphs with important information.</p>
    <footer>Footer content</footer>
    </body>
    </html>
    """
    
    result = extract_web_content(html, "https://example.com")
    print(f"Extracted title: {result.title}")
    print(f"Extracted text ({len(result.text)} chars):")
    print(result.text[:200])
